import io
import tempfile
import typing
from dataclasses import dataclass

import docx
from docx.text import paragraph
from docx.enum.style import WD_STYLE_TYPE

@dataclass
class WordDocument:
    document: docx.Document

    @classmethod
    def from_bytes(cls, contents: bytes):
        doc = docx.Document(io.BytesIO(contents))
        return cls(doc)

    def to_bytes(self) -> bytes:
        with tempfile.TemporaryFile() as tmp_file:
            self.document.save(tmp_file)
            tmp_file.seek(0)
            return tmp_file.read()

    @property
    def paragraphs(self) -> typing.List[paragraph.Paragraph]:
        return self.document.paragraphs

    @property
    def text(self) -> str:
        return "".join(para.text if len(para.text) else " " for para in self.paragraphs)

    def default_paragraph_font(self):
        """ Determines the deafult paragraph font used within the document. """

        if self.doc.document.styles.default(WD_STYLE_TYPE.PARAGRAPH):
            return self.doc.document.styles.default(WD_STYLE_TYPE.PARAGRAPH).font

        return None
