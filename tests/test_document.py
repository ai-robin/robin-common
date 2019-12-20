import io
from unittest.mock import Mock, patch

import docx

from robin_common import WordDocument


def test_from_bytes_to_docx():
    with open("test_data/test_file.docx", "rb") as f:
        contents = f.read()
    doc = WordDocument.from_bytes(contents)

    assert doc.text == "Line 1\n \nLine 2\n \nLine 3"


def test_from_docx_to_bytes():
    # Create a document with python-docx
    doc = docx.Document()
    lines = ["Line 1", "", "Line 2", "", "Line 3"]
    for line in lines:
        doc.add_paragraph(line)

    # Load it into WordDocument and output to bytes
    word_doc = WordDocument(doc)
    output = word_doc.to_bytes()

    # Load these bytes back into python-docx to verify integrity
    doc = docx.Document(io.BytesIO(output))
    paras = [para.text for para in doc.paragraphs]
    assert paras == lines


def test_text_extraction():
    paragraphs = [
        Mock(text="a"),
        Mock(text=""),
        Mock(text=""),
        Mock(text="b"),
        Mock(text=""),
        Mock(text="c"),
        Mock(text="d"),
        Mock(text=""),
    ]
    with patch.object(WordDocument, "paragraphs", new=paragraphs):
        assert WordDocument(Mock).text == "a\n \n \nb\n \nc\nd\n "
